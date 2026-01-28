#!/usr/bin/env python3
import argparse
from pathlib import Path
from email import policy
from email.parser import BytesParser
from email.header import decode_header, make_header

from bs4 import BeautifulSoup

# Optional targets: txt, pdf, docx
# .doc is not implemented because generating true .doc is legacy and brittle

def _decode_header_value(value: str) -> str:
    if not value:
        return ""
    try:
        return str(make_header(decode_header(value)))
    except Exception:
        return value

def extract_eml(path: Path):
    """Return (headers_dict, body_text) for an .eml file."""
    with path.open("rb") as f:
        msg = BytesParser(policy=policy.default).parse(f)

    headers = {
        "Subject": _decode_header_value(msg.get("subject")),
        "From": _decode_header_value(msg.get("from")),
        "To": _decode_header_value(msg.get("to")),
        "Cc": _decode_header_value(msg.get("cc")),
        "Date": _decode_header_value(msg.get("date")),
    }

    text_parts = []
    html_parts = []

    if msg.is_multipart():
        for part in msg.walk():
            if part.get_content_disposition() == "attachment":
                continue  # skip attachments in this simple converter
            ctype = part.get_content_type()
            try:
                if ctype == "text/plain":
                    text_parts.append(part.get_content())
                elif ctype == "text/html":
                    html_parts.append(part.get_content())
            except Exception:
                # Fallback for odd encodings
                payload = part.get_payload(decode=True) or b""
                try:
                    text_parts.append(payload.decode(errors="replace"))
                except Exception:
                    pass
    else:
        ctype = msg.get_content_type()
        try:
            if ctype == "text/plain":
                text_parts.append(msg.get_content())
            elif ctype == "text/html":
                html_parts.append(msg.get_content())
        except Exception:
            payload = msg.get_payload(decode=True) or b""
            text_parts.append(payload.decode(errors="replace"))

    if text_parts:
        body_text = "\n\n".join(text_parts)
    elif html_parts:
        soup = BeautifulSoup("\n\n".join(html_parts), "html.parser")
        body_text = soup.get_text(separator="\n", strip=False)
    else:
        body_text = ""

    return headers, body_text

def as_plaintext(headers, body_text: str) -> str:
    lines = []
    for k in ["Subject", "From", "To", "Cc", "Date"]:
        v = headers.get(k)
        if v:
            lines.append(f"{k}: {v}")
    lines.append("")
    lines.append(body_text or "")
    return "\n".join(lines).rstrip() + "\n"

def save_txt(out_path: Path, headers, body_text: str):
    out_path.write_text(as_plaintext(headers, body_text), encoding="utf-8")

def save_docx(out_path: Path, headers, body_text: str):
    from docx import Document  # lazy import
    doc = Document()
    subject = headers.get("Subject") or "(no subject)"
    doc.add_heading(subject, level=1)
    for k in ["From", "To", "Cc", "Date"]:
        v = headers.get(k)
        if v:
            doc.add_paragraph(f"{k}: {v}")
    doc.add_paragraph("")
    for line in (body_text or "").splitlines():
        doc.add_paragraph(line)
    doc.save(out_path.as_posix())

def save_pdf(out_path: Path, headers, body_text: str):
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Preformatted
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.units import inch

    doc = SimpleDocTemplate(
        out_path.as_posix(),
        pagesize=letter,
        leftMargin=0.75 * inch,
        rightMargin=0.75 * inch,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
    )
    styles = getSampleStyleSheet()
    elems = []
    subject = headers.get("Subject") or "(no subject)"
    elems.append(Paragraph(subject, styles["Heading1"]))
    for k in ["From", "To", "Cc", "Date"]:
        v = headers.get(k)
        if v:
            elems.append(Paragraph(f"<b>{k}:</b> {v}", styles["Normal"]))
    elems.append(Spacer(1, 0.2 * inch))
    code_style = styles["Code"] if "Code" in styles else styles["BodyText"]
    elems.append(Preformatted(body_text or "", code_style))
    doc.build(elems)

def main():
    parser = argparse.ArgumentParser(
        description="Convert all .eml files in the current folder to txt, pdf, or docx. Output goes to ./converted."
    )
    parser.add_argument(
        "-f", "--format", choices=["txt", "pdf", "docx"], required=True,
        help="Target format"
    )
    args = parser.parse_args()

    cwd = Path.cwd()
    out_dir = cwd / "converted"
    out_dir.mkdir(exist_ok=True)

    converted = 0
    failed = 0

    for eml in sorted(cwd.glob("*.eml")):
        base = eml.stem
        out_path = out_dir / f"{base}.{args.format}"
        try:
            headers, body_text = extract_eml(eml)
            if args.format == "txt":
                save_txt(out_path, headers, body_text)
            elif args.format == "docx":
                save_docx(out_path, headers, body_text)
            elif args.format == "pdf":
                save_pdf(out_path, headers, body_text)
            converted += 1
        except Exception as e:
            print(f"Failed to convert {eml.name}: {e}")
            failed += 1

    print(f"Converted {converted} file(s). Failed {failed}. Output: {out_dir}")

if __name__ == "__main__":
    main()
